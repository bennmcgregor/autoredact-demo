from docx import Document
import spacy
from spacy.tokens import DocBin
import re

nlp = spacy.blank("en")
doc_bin = DocBin().from_disk("backend/output.spacy")
docs = list(doc_bin.get_docs(nlp.vocab))
print(len(docs))
# print(docs[0])
# print(len(docs[0]))
# print(docs[0])

# class redactW:
#       def __init__(self, id, span):
#             self.span = span
#             self.docId = id

redact_paras = {}

for i,doc in enumerate(docs):
    if doc.spans:
        # print(doc)
        # for t in doc:
        #       print(t)
        for t in doc.spans['sc']:
            if(not (i in redact_paras)):
               redact_paras[i] = []
            redact_paras[i].append(t)
            # print(doc.spans['sc'])
            # print(i)
            # print(t)
            # print(t.start)
            # print(t.end)
            # print("-------------------")
            # break
#     break
        # print("DEBUGG-------------")
        # print(doc)
#         #print(doc)
#         print(len(doc.spans))

class redact_index:
    def __init__(self, start, end):
        self.start = start
        self.end = end
        self.count = 0

class redact_token:
    def __init__(self, text, is_r):
        self.text = text
        self.is_redacted = is_r

data_doc = Document("backend/static/data.docx")
# print(len(data_doc.paragraphs[49].runs))
for i in redact_paras:
    redact_indices = []
    # print(i)
    para_doc = nlp(data_doc.paragraphs[i].text)
    tokens = [redact_token(x, False) for x in para_doc]
    spaces = [x.whitespace_ for x in para_doc]
    #tokens = re.findall(r"[\w']+|[.,!?;]", data_doc.paragraphs[i].text)
    # print(tokens)
    for span_r in redact_paras[i]:
        start = span_r.start
        while start < span_r.end:
            # tokens[start].text = "X" * len(tokens[start].text)
            tokens[start].is_redacted = True
            start += 1
    output = ""
    for token_i in range(len(tokens)):
        if(tokens[token_i].is_redacted):
            # if(len(output) == 0):
            #     redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
            # else:
            redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
            redact_indices.append(redact_i)
        output += f'{tokens[token_i].text}{spaces[token_i]}'
    # # out_tokens = nlp(output)
    # print(out_tokens)
    out_ic = 0
    recently_redacted = False
    for run_i in range(len(data_doc.paragraphs[i].runs)):
        run_text = data_doc.paragraphs[i].runs[run_i].text
        # run_tokens = nlp(run_text)
        # print(run_tokens)
        data_doc.paragraphs[i].runs[run_i].text = ""
        for rt_i in range(len(run_text)):
            # if out_ic < len(tokens):
            if run_text[rt_i]:
                if len(redact_indices) and redact_indices[0].start <= out_ic and redact_indices[0].end > out_ic:
                    redact_indices[0].count += 1
                    print(redact_indices[0].start, redact_indices[0].end, redact_indices[0].count)
                    print(redact_indices[0].count)
                    print(redact_indices[0].end - redact_indices[0].start)
                    if(redact_indices[0].count == redact_indices[0].end - redact_indices[0].start):
                        print(recently_redacted)
                        if(not recently_redacted):
                            data_doc.paragraphs[i].runs[run_i].text += "XXXX"
                        redact_indices.pop(0)
                        recently_redacted = True
                else:
                    data_doc.paragraphs[i].runs[run_i].text += output[out_ic]
                    recently_redacted = False
                out_ic += 1


    #data_doc.paragraphs[i].text = output
    # print(output)
    # print(nlp(tokens))

data_doc.save('redacted.docx')

# print(data_doc.paragraphs[49].text)
# for i in range(len(data_doc.paragraphs[49].text.split(" "))):
#     print(i)
#     print(data_doc.paragraphs[49].text.split(" ")[i])

# print(len(data_doc.paragraphs))
# paras = []
# for i, p in enumerate(doc.paragraphs):
#     for ir, pr in enumerate(doc.paragraphs):
#     print(p.text)
        # paras.append({
        # "i": i,
        # "text": p.text
        # })





# test_bin = DocBin().from_disk("../data/paired_data/test_data.spacy")
# tests = list(test_bin.get_docs(nlp.vocab))
# print("HIII")
# for i,test in enumerate(tests):
#         #print(test)
#         print(test.spans["sc"])
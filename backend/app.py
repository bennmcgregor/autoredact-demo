import json
from flask import Flask, request, render_template
from docx import Document
import subprocess
from docx import Document
import spacy
from spacy.tokens import DocBin, SpanGroup, Span
from docx2pdf import convert
from flask_cors import CORS
import pandas as pd
import plotly
import plotly.express as px
from relabel_paired_data import relabelled_doc
from redact_doc import redact_doc

app = Flask(__name__)
CORS(app)

# test_bin = DocBin().from_disk("../data/paired_data/test_data.spacy")
# tests = list(test_bin.get_docs(nlp.vocab))
# #print("HIII")
# for i,test in enumerate(tests):
#         ##print(test)
#         #print(test.spans["sc"])

def extract_paras(r, u):
    r_nonempty_paras = []
    u_nonempty_paras = []
    for para in r.paragraphs:
        stripped = para.text.strip()
        if stripped != "":
            r_nonempty_paras.append(" ".join(stripped.split()))
   
    for para in u.paragraphs:
      stripped = para.text.strip()
      if stripped != "":
            u_nonempty_paras.append(" ".join(stripped.split())) 
    
    return r_nonempty_paras, u_nonempty_paras

def get_unredacted_spans(r_doc):
    unredacted_spans = []
    end_idx = 0
    while True:
        start_idx = end_idx

        while True:
            token = r_doc[end_idx]
            if token.text.startswith('XX'):
                break
            end_idx += 1
            if end_idx == len(r_doc):
                unredacted_spans.append(r_doc[start_idx:])
                return unredacted_spans
        if end_idx > start_idx:
            unredacted_spans.append(r_doc[start_idx:end_idx])
            start_idx = end_idx
        while True:
            end_idx += 1
            if end_idx == len(r_doc):
                return unredacted_spans
            token = r_doc[end_idx]
            if not token.text.startswith('XX'):
                break

def verify_matching_section(start_idx, span, u_doc):
    for sec_tk, para_tk in zip(span, u_doc[start_idx:start_idx + len(span)]):
        if sec_tk.text != para_tk.text:
            return False
    return True

def find_first_matching_section(start_idx, span, u_doc):
    for i in range(start_idx, len(u_doc)):
        if u_doc[i].text == span[0].text and verify_matching_section(i, span, u_doc):
            return i

def annotate_paragraph(u_doc, r_doc):
    unredacted_spans = get_unredacted_spans(r_doc)

    match_idx = 0
    prev_idx = 0
    redacted_spans = list()
    for span in unredacted_spans:
        match_idx = find_first_matching_section(prev_idx, span, u_doc)
        if match_idx == None:
            #print("COULD NOT FIND MATCH:")
            #print(u_doc)
            #print("---------")
            #print(span)
            #print("xxxxxxxxxx\n")
            raise Exception("Could not find match")
        if match_idx > prev_idx:
            redacted_spans.append(Span(u_doc, prev_idx, match_idx, "REDACTED"))
        prev_idx = min(len(u_doc), match_idx + len(span))
    
    if prev_idx < len(u_doc):
        redacted_spans.append(Span(u_doc, prev_idx, len(u_doc), "REDACTED"))

    u_doc.spans["sc"] = SpanGroup(u_doc, spans=redacted_spans)

    # if len(redacted_spans) > 0:
    #     #print(u_doc.spans)

    return u_doc

def read_paired_data(nlp, filename):
    with open(filename, "r", encoding="utf8") as f:
        for line in f:
            #print("LINE", line)
            yield nlp(json.loads(line)["unredacted"]), nlp(json.loads(line)["redacted"])


@app.route('/')
def home_page():
   return render_template("index.html")

@app.route('/upload', methods = ['POST'])
def upload_file():
    if request.method == 'POST':
      filename = request.form['filename']
      #f = request.files['file']
      #if not os.path.isfile('./pdfs/'+filename):
          
      convert('./unredacted/'+filename, '../frontend_2/src/data.pdf' )

      #f.save('../frontend_2/src/data.docx')
      #convert('../frontend_2/src/data.docx', '../frontend_2/src/data.pdf')

      return "success", 200

@app.route('/redactor', methods = ['POST'])
def redact():
   if request.method == 'POST':
      #f = request.files['file']
      filename = request.form['filename']

      doc = Document("./unredacted/"+filename)
      paras = []

      for i, p in enumerate(doc.paragraphs):
          paras.append({
             "i": i,
             "text": p.text
          })

      for i,t in enumerate(doc.tables):
          for r in t.rows:
              for c in r.cells:
                  for para in c.paragraphs:
                     paras.append({
                        "t": i,
                        "text": para.text
                     })


      with open("data.jsonl", "w", encoding="utf-8") as output:
        while len(paras) > 0:
            para = paras.pop(0)
            output.write(json.dumps(para) + "\n")
      
      subprocess.run(["python", "-m", "spacy", "apply", "./models/en_trf_docs_v5_bs_2000_cont/model-best", "data.jsonl", "output.spacy", "--force"])

      redact_doc(filename)

      return "success", 200
   



def figures_to_html(figs):
   #  with open(filename, 'w', encoding="utf-8") as dashboard:
      #   dashboard.write("<html><head></head><body>" + "\n")
      html = []
      for fig in figs:
         inner_html = fig.to_html().split('<body>')[1].split('</body>')[0]
         html.append(inner_html)
      return html
      #   dashboard.write("</body></html>" + "\n")

@app.route('/graphs', methods = ['POST', 'GET'])
def graphs():
   df = pd.read_csv('Training Results - Sheet1.csv')
   df.drop(df.index[0])
   df.columns = df.iloc[0]
   df_labels = df.copy()


   df_labels.dropna(subset=['Label', 'Name'], inplace=True)

   new_table = pd.DataFrame(columns=['Model Name', 'Label', 'Score', 'Score Type', 'Params', 'Name'])
   row_id = 0
   for index, row in df_labels.iloc[:,0:7].iterrows():
      if(index != 0):
         new_table.loc[row_id] = [row['Model Name'], row['Label'], row['P'], 'Precision', row["Parameters (differerent from default)"], row["Name"]]
         row_id += 1
         new_table.loc[row_id] = [row['Model Name'], row['Label'], row['R'], 'Recall', row["Parameters (differerent from default)"], row["Name"]]
         row_id += 1
         new_table.loc[row_id] = [row['Model Name'], row['Label'], row['F'], 'F1 Score', row["Parameters (differerent from default)"], row["Name"]]
         row_id += 1
   new_table.fillna('', inplace=True)

   labels = []
   color_discrete_sequence = []
   for i in range(len(new_table)//3//3):
      labels.extend(["REDACTED", "ID", "CONTEXT"])
      color_discrete_sequence = ["#D8FAFD", '#3E5C76', '#5c6f87']

   #print(new_table)
   new_table["Score"] = new_table["Score"].astype(float)
   #print(new_table.dtypes)

   model2_df = new_table[(new_table["Name"] == "RoBERTa-DOCS_V5") | (new_table["Name"] == "RoBERTa-V5")]
   #print(model2_df)

   labels_m2 = []
   color_discrete_sequence = []
   for i in range(len(model2_df)//3):
      if i % 3 == 0:
         labels_m2.extend(["REDACTED", "REDACTED", "REDACTED"])
      if i % 3 == 1:
           labels_m2.extend(["ID", "ID", "ID"])
      if i % 3 == 1:
         labels_m2.extend(["CONTEXT", "CONTEXT", "CONTEXT"])
      #color_discrete_sequence = ["#D8FAFD", '#3E5C76', '#5c6f87']
      color_discrete_sequence = ["#ABD8EF", '#EDF7F6', '#5c6f87']
   #print(labels_m2)

   modelb_df = model2_df[(new_table["Name"] == "RoBERTa-DOCS_V5")]
   #print(modelb_df)

   model2_fig = px.bar(modelb_df.loc[new_table['Label'] == "REDACTED"], x="Score Type", y="Score",
   color="Name", barmode = 'group', hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="Overall Performance of RoBERTa-DOCS_v5", text="Score")
   model2_fig.update_traces(texttemplate='%{value:.3g}')
   model2_fig.update_layout(title={'font': {'size': 12}}, yaxis_title="Percentage (%)", xaxis_title=None, font=dict(size=6))
   model2_fig.update_layout(showlegend=False)
   model2_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   model2_labels_fig = px.bar(model2_df, x="Score Type", y="Score",
   color="Name", barmode = 'group', hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="Model Performance by Redaction Type")
   model2_labels_fig.update_traces(text=labels_m2, insidetextanchor = "middle", textangle=0)
   model2_labels_fig.update_layout(title={'font': {'size': 12}}, 
   yaxis_title="Percentage (%)", xaxis_title=None, legend={"y":0.5},  font=dict(size=6))
   model2_labels_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   modelb_labels_fig = px.bar(modelb_df.loc[new_table['Label'] != "REDACTED"], x="Label", y="Score",
   color="Score Type", barmode = 'group', hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="RoBERTa-DOCS_V5 Performance by Redaction Type", facet_col="Score Type", text="Score")
   modelb_labels_fig.update_layout(title={'font': {'size': 12}}, font=dict(size=6), yaxis_title="Percentage (%)", xaxis_title=None, legend={"y":0.5})
   modelb_labels_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   modelb_labels_fig_alt = px.bar(modelb_df.loc[new_table['Label'] != "REDACTED"], x="Score Type", y="Score",
   color="Label", barmode = 'group', hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="RoBERTa-DOCS_V5 Performance by Redaction Type", text="Score")
   modelb_labels_fig_alt.update_layout(title={'font': {'size': 12}}, font=dict(size=6), yaxis_title="Percentage (%)", xaxis_title=None, legend={"y":0.5})
   modelb_labels_fig_alt.update_layout(hoverlabel = dict(font=dict(size=8)))


   f3_group = ["RoBERTa-base", "RoBERTa-V2", "RoBERTa-V3", "RoBERTa-V4", "RoBERTa-DOCS_V5", "RoBERTa-V5"]
   f4_group = ["RoBERTa-base", "spaCy", "SpanBERT", "Legal-BERT", "RoBERTa-DOCS_V5", "RoBERTa-V5"]
   f5_group = ["GPT-3", "naive-1000", "presidio", "RoBERTa-DOCS_V5", "RoBERTa-V5"]

   model_f3_df = new_table[(new_table["Label"] == "REDACTED") & (new_table["Score Type"] == "F1 Score") & (new_table["Name"].isin(f3_group))]
   #print(model_f3_df)

   model_f4_df = new_table[(new_table["Label"] == "REDACTED") & (new_table["Score Type"] == "F1 Score") & (new_table["Name"].isin(f4_group))]
   #print(model_f4_df)

   model_f5_df = new_table[(new_table["Label"] == "REDACTED") & (new_table["Score Type"] == "F1 Score") & (new_table["Name"].isin(f5_group))]
   #print(model_f5_df)

   f3g_fig = px.bar(model_f3_df, x="Name", y="Score", color="Name",  hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=['#EDF7F6', "#ABD8EF", '#5c6f87', "#ABD8EF", '#EDF7F6', '#5c6f87'], title="Model Improvement Over Time", text="Score")
   f3g_fig.update_layout(showlegend=False)
   f3g_fig.update_layout(xaxis={'categoryorder':'total ascending'})
   f3g_fig.update_traces(texttemplate='%{value:.3g}')
   f3g_fig.update_layout(title={'font': {'size': 12}}, yaxis_title="F1 Score (%)", xaxis_title="Model Version", font=dict(size=6))
   f3g_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   f4g_fig = px.bar(model_f4_df, x="Name", y="Score", color="Name",  hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=['#EDF7F6', "#ABD8EF", '#5c6f87', "#ABD8EF", '#EDF7F6', '#5c6f87'], title="Comparison with Alternative Transformer Models", text="Score")
   f4g_fig.update_layout(showlegend=False)
   f4g_fig.update_layout(xaxis={'categoryorder':'total ascending'})
   f4g_fig.update_traces(texttemplate='%{value:.3g}')
   f4g_fig.update_layout(title={'font': {'size': 12}}, yaxis_title="F1 Score (%)", xaxis_title="Model Version Number", font=dict(size=6))
   f4g_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   f5g_fig = px.bar(model_f5_df, x="Name", y="Score", color="Name",  hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="Comparison with Alternative Approaches & Baseline", text="Score")
   f5g_fig.update_layout(showlegend=False)
   f5g_fig.update_layout(xaxis={'categoryorder':'total ascending'})
   f5g_fig.update_traces(texttemplate='%{value:.3g}')
   f5g_fig.update_layout(title={'font': {'size': 12}}, yaxis_title="F1 Score (%)", xaxis_title=None, font=dict(size=6))
   f5g_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   model_geo_bias_df = new_table[(new_table["Model Name"] == "Geography bias test para 0") & (new_table["Name"] == "AVG") & (new_table["Score Type"] == "F1 Score")]
   #print(model_geo_bias_df)
   geo_bias_fig = px.bar(model_geo_bias_df, x="Label", y="Score", color="Label",  hover_data=["Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="Performance on Racial Names by Country", text="Score")
   geo_bias_fig.update_layout(showlegend=False)
   geo_bias_fig.update_layout(xaxis={'categoryorder':'total ascending'},  xaxis_title="Country of Origin")
   geo_bias_fig.update_traces(texttemplate='%{value:.3g}')
   geo_bias_fig.update_layout(title={'font': {'size': 12}}, yaxis_title="F1 Score (%)", xaxis_title="Country of Origin", font=dict(size=6))
   geo_bias_fig.update_layout(hoverlabel = dict(font=dict(size=8)))

   # fig1 = px.bar(new_table.loc[new_table['Label'] == "REDACTED"], x="Model Name", y="Score",
   #           color="Score Type", barmode = 'group', hover_data=["Model Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="P, R, and F Scores for REDACTED")

   # fig2 = px.bar(new_table.loc[new_table['Label'] == "ID"], x="Model Name", y="Score",
   #             color="Score Type", barmode = 'group', hover_data=["Model Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence,  title="P, R, and F Scores for ID")

   # fig3 = px.bar(new_table.loc[new_table['Label'] == "CONTEXT"], x="Model Name", y="Score",
   #             color="Score Type", barmode = 'group', hover_data=["Model Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence,  title="P, R, and F Scores for CONTEXT")

   # fig4 = px.bar(new_table, x="Model Name", y="Score",
   #             color="Score Type", barmode = 'group', facet_col="Label",hover_data=["Model Name", "Score", "Score Type", "Label", "Params"], color_discrete_sequence=color_discrete_sequence, title="P, R, and F Scores for REDACTED, ID, and CONTEXT")

   # fig5 = px.bar(new_table, x="Model Name", y="Score",
   #           color="Score Type", barmode = 'group', hover_data=["Model Name", "Score", "Score Type", "Label", "Params"],color_discrete_sequence=color_discrete_sequence,  title="P, R, and F Scores for REDACTED, ID, and CONTEXT" )

   # fig5.update_traces(text=labels, insidetextanchor = "middle", textangle=0)

   plot_json = []
   plot_json.append(plotly.io.to_json(model2_fig,pretty=False))
   plot_json.append(plotly.io.to_json(model2_labels_fig,pretty=False))
   plot_json.append(plotly.io.to_json(modelb_labels_fig,pretty=False))
   plot_json.append(plotly.io.to_json(modelb_labels_fig_alt,pretty=False))
   plot_json.append(plotly.io.to_json(f3g_fig,pretty=False))
   plot_json.append(plotly.io.to_json(f4g_fig,pretty=False))
   plot_json.append(plotly.io.to_json(f5g_fig,pretty=False))
   plot_json.append(plotly.io.to_json(geo_bias_fig,pretty=False))

   # plot_json.append(plotly.io.to_json(fig5, pretty=False))
   # plot_json.append(plotly.io.to_json(fig4, pretty=False))
   # plot_json.append(plotly.io.to_json(fig1, pretty=False))
   # plot_json.append(plotly.io.to_json(fig2, pretty=False))
   # plot_json.append(plotly.io.to_json(fig3, pretty=False))

   # plot_html = figures_to_html([fig4, fig5])

   return json.dumps(plot_json)
         
@app.route('/getstats', methods = ['POST'])
def get_stats():
      if request.method == 'POST':
         filename = request.form['filename']

         #doc = Document('./unredacted/' + filename)
         #ground_truth = Document('./ground_truths/' + filename)

         # non empty paras in (ground truth) redacted and unredacted docx files
         #r_paras, u_paras = extract_paras(ground_truth, doc)
         #print("check", u_paras)

         # Create dataframe and jsonl file
         #df = pd.DataFrame(zip(u_paras, r_paras), columns=['unredacted', 'redacted'])
         #df.to_json('paired_data.jsonl', orient='records', lines=True)

         # Create spacy file
         nlp = spacy.blank("en")
         doc_bin = DocBin(attrs=[])
         file_path = "./ground_jsonl/" + filename.split(".")[0]+".jsonl"
         for u_para, r_para in read_paired_data(nlp, file_path):
            # #print("*******************************************")
            try:
               doc_bin.add(annotate_paragraph(u_para, r_para))
            except:
               #print("EXCEPTION!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
               continue
         doc_bin.to_disk("paired_data.spacy")

      # Relabel paired data
      relabelled_doc()

      # Evaluate
      output = subprocess.check_output(["python", "-m", "spacy", "evaluate", "./models/en_trf_docs_v5_bs_2000_cont/model-best", "relabelled_paired_data.spacy", "--output", "eval.json"]).decode("utf-8")
      # parse the output and return it to the front end
      output = output.split("============================== SPANS (per type) ==============================", 1)[1]
      output = output.split()[4:16]
      json_resp = {
            output[0]: {
                  "P": output[1],
                  "R": output[2],
                  "F": output[3],
            },
            output[4]: {
               "P": output[5],
               "R": output[6],
               "F": output[7],
            },
            output[8]: {
               "P": output[9],
               "R": output[10],
               "F": output[11],
            },  
      }
      #print(json_resp)

      return json_resp
		
if __name__ == '__main__':
   app.run(debug = False, threaded=True)
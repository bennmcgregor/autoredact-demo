from docx import Document
import spacy
import subprocess
from spacy.tokens import DocBin, SpanGroup, Span
from docx2pdf import convert

class redact_index:
    def __init__(self, start, end):
        self.start = start
        self.end = end
        self.count = 0

class redact_token:
    def __init__(self, text, is_r):
        self.text = text
        self.is_redacted = is_r

def redact_doc():
   nlp = spacy.blank("en")
   doc_bin = DocBin().from_disk("output.spacy")
   docs = list(doc_bin.get_docs(nlp.vocab))
   data_doc = Document("../frontend_2/src/data.docx")

   redact_paras = {}
   redact_cells = {}

   num_paras = len(data_doc.paragraphs)
   for i,doc in enumerate(docs):
      if doc.spans:
         for t in doc.spans['sc']:
               if i < num_paras:
                  if(not (i in redact_paras)):
                     redact_paras[i] = []
                  redact_paras[i].append(t)
               else:
                  if(not (i in redact_cells)):
                     redact_cells[i-num_paras] = []
                  redact_cells[i-num_paras].append(t)
   data_doc_cells = []
   for i,t in enumerate(data_doc.tables):
          for r in t.rows:
              for c in r.cells:
                  for para in c.paragraphs:
                     data_doc_cells.append(para)

   for i in redact_paras:
      redact_indices = []
      para_doc = nlp(data_doc.paragraphs[i].text)
      tokens = [redact_token(x, False) for x in para_doc]
      spaces = [x.whitespace_ for x in para_doc]
      #tokens = re.findall(r"[\w']+|[.,!?;]", data_doc.paragraphs[i].text)
      # print(tokens)
      for span_r in redact_paras[i]:
         start = span_r.start
         while start < span_r.end:
               # tokens[start].text = "X" * len(tokens[start].text)
               tokens[start].is_redacted = True
               start += 1
      output = ""
      for token_i in range(len(tokens)):
         if(tokens[token_i].is_redacted):
               # if(len(output) == 0):
               #     redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
               # else:
               redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
               redact_indices.append(redact_i)
         output += f'{tokens[token_i].text}{spaces[token_i]}'
      # # out_tokens = nlp(output)
      # print(out_tokens)
      out_ic = 0
      recently_redacted = False
      print(len(data_doc.paragraphs[i].runs))
      data_doc_runs = data_doc.paragraphs[i].runs
      data_doc.paragraphs[i].clear()
      for run_i in range(len(data_doc_runs)):
         run_text = data_doc_runs[run_i].text
         #data_doc.paragraphs[i].runs[run_i].text = ""
         curr = data_doc.paragraphs[i].add_run(text="",style=data_doc_runs[run_i].style)
         for rt_i in range(len(run_text)):
               # if out_ic < len(tokens):
               if run_text[rt_i]:
                  if len(redact_indices) and redact_indices[0].start <= out_ic and redact_indices[0].end > out_ic:
                     redact_indices[0].count += 1
                     if(redact_indices[0].count == redact_indices[0].end - redact_indices[0].start):
                           if(not recently_redacted):
                              run = data_doc.paragraphs[i].add_run(text='XXXX')
                              run.font.highlight_color = 1
                              curr = data_doc.paragraphs[i].add_run(text="",style=data_doc_runs[run_i].style)
                              #data_doc.paragraphs[i].runs[run_i].text += "XXXX"
                           redact_indices.pop(0)
                           recently_redacted = True
                  else:
                     curr.text += output[out_ic]
                     recently_redacted = False
                  out_ic += 1

   for i in redact_cells:
      redact_indices = []
      print("cell",i)
      para_doc = nlp(data_doc_cells[i].text)
      tokens = [redact_token(x, False) for x in para_doc]
      spaces = [x.whitespace_ for x in para_doc]
      #tokens = re.findall(r"[\w']+|[.,!?;]", data_doc.paragraphs[i].text)
      # print(tokens)
      for span_r in redact_cells[i]:
         start = span_r.start
         while start < span_r.end:
               # tokens[start].text = "X" * len(tokens[start].text)
               tokens[start].is_redacted = True
               start += 1
      output = ""
      for token_i in range(len(tokens)):
         if(tokens[token_i].is_redacted):
               # if(len(output) == 0):
               #     redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
               # else:
               redact_i = redact_index(len(output), len(output) + len(tokens[token_i].text))
               redact_indices.append(redact_i)
         output += f'{tokens[token_i].text}{spaces[token_i]}'
      # # out_tokens = nlp(output)
      # print(out_tokens)
      out_ic = 0
      recently_redacted = False

      data_doc_cell_runs = data_doc_cells[i].runs
      data_doc_cells[i].clear()
      for run_i in range(len(data_doc_cell_runs)):
         run_text = data_doc_cell_runs[run_i].text
         #data_doc_cells[i].runs[run_i].text = ""
         curr = data_doc_cells[i].add_run(text="",style=data_doc_runs[run_i].style)
         for rt_i in range(len(run_text)):
               # if out_ic < len(tokens):
               if run_text[rt_i]:
                  if len(redact_indices) and redact_indices[0].start <= out_ic and redact_indices[0].end > out_ic:
                     redact_indices[0].count += 1
                     if(redact_indices[0].count == redact_indices[0].end - redact_indices[0].start):
                           if(not recently_redacted):
                              run = data_doc_cells[i].add_run(text='XXXX')
                              run.font.highlight_color = 1
                              curr = data_doc_cells[i].add_run(text="",style=data_doc_runs[run_i].style)
                              #data_doc_cells[i].runs[run_i].text += "XXXX"
                           redact_indices.pop(0)
                           recently_redacted = True
                  else:
                     curr.text += output[out_ic]
                     recently_redacted = False
                  out_ic += 1

   data_doc.save('../frontend_2/src/redacted.docx')
   convert("../frontend_2/src/redacted.docx", "../frontend_2/src/redacted.pdf")
